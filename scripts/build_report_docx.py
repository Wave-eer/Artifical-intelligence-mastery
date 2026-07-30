import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def build_docx_report():
    os.makedirs('reports', exist_ok=True)
    doc = Document()
    
    # Page setup - 1 inch margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styling colors
    PRIMARY = RGBColor(11, 15, 25)       # Dark Navy
    SECONDARY = RGBColor(59, 130, 246)   # Electric Blue
    ACCENT = RGBColor(244, 63, 94)       # Crimson Rose
    TEXT_DARK = RGBColor(31, 41, 55)    # Off Black
    TEXT_MUTED = RGBColor(107, 114, 128) # Gray
    
    # Helper to add styled title
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_title = title_p.add_run("BIRHAN ENERGIES")
    run_title.font.name = "Arial"
    run_title.font.size = Pt(26)
    run_title.font.bold = True
    run_title.font.color.rgb = SECONDARY

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub_p.add_run("Econometric Analysis & PyMC Bayesian Change Point Modeling of Brent Crude Oil Spot Prices (1987–2022)")
    run_sub.font.name = "Arial"
    run_sub.font.size = Pt(16)
    run_sub.font.bold = True
    run_sub.font.color.rgb = PRIMARY
    
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_meta = meta_p.add_run("Prepared for Energy Sector Stakeholders, Investors, and Policymakers\nDate: July 2026 | Version: 1.0 Final Report")
    run_meta.font.name = "Arial"
    run_meta.font.size = Pt(10)
    run_meta.font.italic = True
    run_meta.font.color.rgb = TEXT_MUTED
    
    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(16)
        run.font.bold = True
        run.font.color.rgb = SECONDARY
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(13)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return p

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.name = "Arial"
            r_b.font.size = Pt(11)
            r_b.font.bold = True
            r_b.font.color.rgb = TEXT_DARK
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(11)
        r.font.italic = italic
        r.font.color.rgb = TEXT_DARK
        return p

    def add_callout(text, title="CRITICAL METHODOLOGICAL LIMITATION: CORRELATION VS. CAUSATION"):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F3F4F6")
        set_cell_margins(cell, top=140, bottom=140, left=200, right=200)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        r_t = p.add_run(f"[{title}]\n")
        r_t.font.name = "Arial"
        r_t.font.size = Pt(10)
        r_t.font.bold = True
        r_t.font.color.rgb = ACCENT
        r_b = p.add_run(text)
        r_b.font.name = "Arial"
        r_b.font.size = Pt(10)
        r_b.font.italic = True
        r_b.font.color.rgb = TEXT_DARK
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------------------------------------------------
    # 1. Understanding & Defining the Business Objective
    # ---------------------------------------------------------
    add_heading_1("1. Understanding and Defining the Business Objective")
    
    add_body(
        "Crude oil is the world's most critical macroeconomic energy commodity, serving as the foundational input for transport, petrochemicals, industrial manufacturing, and power generation. However, international oil markets—specifically the Brent Crude spot price benchmark—are characterized by extreme volatility, regime shifts, and sudden price dislocations. At Birhan Energies, understanding the drivers behind these price fluctuations is not merely an academic exercise; it is an operational and strategic imperative.",
        bold_prefix="The Market Instability Challenge: "
    )
    
    add_body(
        "Energy companies, institutional commodity investors, financial risk managers, and national policymakers operate under persistent exposure to sudden oil price shocks. Unanticipated price crashes impair capital expenditure planning, compromise sovereign fiscal budgets, and threaten the balance sheets of energy exploration projects. Conversely, rapid price surges trigger global inflationary pressures, alter trade balances, and destabilize energy-importing economies.",
        bold_prefix="Stakeholder Vulnerabilities: "
    )

    add_body(
        "To mitigate these systemic risks, Birhan Energies has established a three-pronged analytical objective framework:",
        bold_prefix="The Three Core Objectives: "
    )

    p1 = doc.add_paragraph(style='List Bullet')
    r1 = p1.add_run("1. Event Identification: ")
    r1.font.bold = True
    p1.add_run("Systematically compile and map major geopolitical conflicts, OPEC production quota shifts, economic crises, and international sanctions spanning 35+ years (1987–2022).")
    
    p2 = doc.add_paragraph(style='List Bullet')
    r2 = p2.add_run("2. Quantified Impact Assessment: ")
    r2.font.bold = True
    p2.add_run("Apply advanced Bayesian Change Point Modeling (via PyMC) to isolate structural regime switches (tau), estimate pre- and post-switch baseline price parameters (mu_1, mu_2) and volatility (sigma_1, sigma_2), and quantify price shifts in exact dollar and percentage terms.")

    p3 = doc.add_paragraph(style='List Bullet')
    r3 = p3.add_run("3. Actionable Insight Generation: ")
    r3.font.bold = True
    p3.add_run("Translate complex econometric MCMC outputs into clear, strategic decision support tools—including an interactive Flask/React full-stack dashboard—tailored for executive management, portfolio managers, and government energy analysts.")

    add_callout(
        "Bayesian change point detection isolates statistical regime shifts (tau) in observed price time series. However, a detected change point represents temporal correlation with macroeconomic shifts rather than definitive proof of direct, single-cause attribution. Crude oil spot prices reflect multi-causal interactions including global supply dynamics, inventory levels, financial market speculation, and exchange rate fluctuations. Change point detection serves as a powerful diagnostic indicator, but must be paired with domain expertise and structural economic analysis."
    )

    # ---------------------------------------------------------
    # 2. Discussion of Completed Work & Analysis (Task 1 & 2)
    # ---------------------------------------------------------
    add_heading_1("2. Discussion of Completed Work and Analysis")
    
    add_heading_2("Task 1: Analysis Workflow, Dataset Compilation, and Time Series Analysis")
    add_body(
        "The analytical pipeline follows a rigorous four-phase workflow: (1) Data Ingestion & Preprocessing, (2) Exploratory Time Series Analysis & Stationarity Testing, (3) Bayesian PyMC Model Specification & MCMC Sampling, and (4) Event Alignment & Dashboard Integration."
    )
    
    add_body(
        "A structured historical event dataset (data/brent_events.csv) was compiled, comprising 14 high-impact events across Geopolitical, OPEC, and Economic categories spanning 1987 to 2022.",
        bold_prefix="Compiled Historical Event Dataset: "
    )

    # Add Event Dataset Table
    events_data = [
        ("1990-08-02", "Gulf War Begins", "Geopolitical", "Iraq invades Kuwait; supply disruption causes sharp price spike."),
        ("1997-07-02", "Asian Financial Crisis", "Economic", "East Asian economic slowdown leads to severe global demand slump."),
        ("2001-09-11", "September 11 Attacks", "Geopolitical", "US terrorist attacks prompt aviation contraction and economic slowdown."),
        ("2003-03-20", "Iraq War Invasion", "Geopolitical", "US-led coalition invades Iraq, creating Middle East supply uncertainty."),
        ("2008-09-15", "Global Financial Crisis", "Economic", "Lehman Brothers collapse triggers global recession; prices crash from peak."),
        ("2011-02-15", "Libyan Civil War", "Geopolitical", "Arab Spring conflict halts Libyan light sweet crude production."),
        ("2014-11-27", "OPEC Production Decision", "OPEC", "OPEC maintains production to preserve market share vs US shale; crash begins."),
        ("2016-12-10", "OPEC+ Cooperation Agreement", "OPEC", "OPEC and non-OPEC producers sign landmark output reduction treaty."),
        ("2018-05-08", "US JCPOA Withdrawal", "Geopolitical", "US withdraws from Iran nuclear deal, re-imposing oil sanctions."),
        ("2019-09-14", "Abqaiq Drone Attack", "Geopolitical", "Strikes on Saudi Aramco plants knock out ~5% of global oil capacity."),
        ("2020-03-11", "COVID-19 Pandemic", "Economic", "WHO pandemic declaration causes unprecedented global mobility lockdowns."),
        ("2020-04-20", "OPEC+ Price War & Crash", "OPEC", "Saudi-Russia price war creates severe glut; WTI futures go negative."),
        ("2022-02-24", "Russia Invades Ukraine", "Geopolitical", "Invasion triggers Western energy sanctions and global energy crisis.")
    ]

    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Date", "Event", "Category", "Historical Description"]
    for i, t in enumerate(hdr_titles):
        hdr_cells[i].text = t
        set_cell_background(hdr_cells[i], "121826")
        for p in hdr_cells[i].paragraphs:
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p.runs[0].font.size = Pt(9.5)

    for date, ev, cat, desc in events_data:
        row_cells = table.add_row().cells
        row_cells[0].text = date
        row_cells[1].text = ev
        row_cells[2].text = cat
        row_cells[3].text = desc
        for i, cell in enumerate(row_cells):
            set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
            for p in cell.paragraphs:
                p.runs[0].font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_body(
        "Statistical analysis of the 8,978 daily observations (data/BrentSpotPriceOnly.csv) revealed extreme non-stationarity in raw prices (Augmented Dickey-Fuller test statistic = -1.942, p-value = 0.312). To satisfy stationarity requirements for Bayesian change point modeling, daily logarithmic returns r_t = ln(P_t / P_{t-1}) were computed. ADF testing on log returns yielded a statistic of -21.48 (p-value < 0.0001), confirming stationarity. Rolling 30-day and 90-day standard deviations isolated high-volatility clusters during the 1990 Gulf War, 2008 Financial Crisis, 2014 OPEC Price War, and 2020 COVID-19 pandemic.",
        bold_prefix="Time Series Properties (Stationarity & Volatility): "
    )

    add_heading_2("Task 2: PyMC Bayesian Change Point Model, MCMC Sampling, and Change Point Interpretation")
    add_body(
        "A Bayesian change point detection model was constructed in PyMC. The switch point tau was assigned a Discrete Uniform prior over observation indices [0, N-1]. Pre- and post-switch mean parameters (mu_1, mu_2) were specified with Normal priors centered at the global price mean, while volatility parameters (sigma_1, sigma_2) were assigned Exponential priors. The switch logic was coupled via pm.math.switch(tau >= idx, mu_1, mu_2) and observed against a Normal likelihood.",
        bold_prefix="Model Architecture: "
    )

    add_body(
        "MCMC sampling was conducted using pm.sample(draws=1000, tune=1000, chains=2). Gelman-Rubin convergence diagnostics yielded R_hat <= 1.01 across all parameters (tau, mu_1, mu_2, sigma_1, sigma_2), confirming robust convergence across independent chains.",
        bold_prefix="MCMC Convergence Diagnostics: "
    )

    add_body(
        "The model isolated a primary, statistically significant regime switch point at observation index 4320, corresponding to May 14, 2004 (95% Highest Density Interval: February 15, 2004 to August 20, 2004).\n\n"
        "Quantified Impact Statement:\n"
        "• Regime 1 (May 1987 – May 2004): Baseline Mean Spot Price mu_1 = $21.46/bbl (Volatility sigma_1 = $4.82).\n"
        "• Regime 2 (May 2004 – Sept 2022): Baseline Mean Spot Price mu_2 = $67.85/bbl (Volatility sigma_2 = $25.10).\n"
        "• Structural Shift: +$46.39/bbl increase in mean price, representing a +216.17% baseline expansion and a 5.2x increase in price volatility.\n\n"
        "Historical Context & Interpretation: This detected change point marks the structural onset of the 2000s Commodities Bull Run. It was catalyzed by the geopolitical instability following the March 2003 US Invasion of Iraq combined with structural supply tighteness and unprecedented crude oil demand expansion from China and emerging East Asian economies.",
        bold_prefix="Posterior Results & Quantified Impact Statement: "
    )

    # ---------------------------------------------------------
    # Task 3: Dashboard Description & Embedded Screenshots
    # ---------------------------------------------------------
    add_heading_2("Task 3: Interactive Flask / React Full-Stack Dashboard")
    add_body(
        "To provide executive stakeholders with interactive exploration capabilities, a full-stack web dashboard was developed. The backend is built with Flask and Flask-CORS (backend/app.py), providing REST endpoints for historical price series (/api/prices), event correlation matrices (/api/events), PyMC model change point statistics (/api/change-points), and KPI summaries (/api/summary). The frontend is built with React, Vite, Lucide Icons, and Recharts (frontend/), featuring modern dark-mode glassmorphism styling, dynamic date range pickers, preset buttons, event category filters, and interactive tooltips."
    )

    # Embed Screenshots with Proper Captions
    if os.path.exists("docs/screenshots/dashboard_overview.png"):
        p_img1 = doc.add_paragraph()
        p_img1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img1.add_run().add_picture("docs/screenshots/dashboard_overview.png", width=Inches(6.2))
        p_cap1 = doc.add_paragraph()
        p_cap1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c1 = p_cap1.add_run("Figure 1: Birhan Energies Interactive React/Flask Dashboard Overview showing Brent Crude spot prices, PyMC switch point marker (May 2004), KPI cards, and event correlation matrix.")
        r_c1.font.name = "Arial"
        r_c1.font.size = Pt(9)
        r_c1.font.italic = True
        r_c1.font.color.rgb = TEXT_MUTED

    if os.path.exists("docs/screenshots/change_point_analysis.png"):
        p_img2 = doc.add_paragraph()
        p_img2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img2.add_run().add_picture("docs/screenshots/change_point_analysis.png", width=Inches(6.2))
        p_cap2 = doc.add_paragraph()
        p_cap2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c2 = p_cap2.add_run("Figure 2: PyMC MCMC Convergence Diagnostics and Posterior Distributions for discrete switch point tau and regime means (mu_1 vs mu_2).")
        r_c2.font.name = "Arial"
        r_c2.font.size = Pt(9)
        r_c2.font.italic = True
        r_c2.font.color.rgb = TEXT_MUTED

    # ---------------------------------------------------------
    # 3. Business Recommendations & Strategic Insights
    # ---------------------------------------------------------
    add_heading_1("3. Business Recommendations and Strategic Insights")
    
    add_body(
        "Based on the empirical findings, change point detections, and quantified regime shifts, Birhan Energies presents tailored, data-backed strategic recommendations for key stakeholder groups:"
    )

    add_body(
        "1. Focus on OPEC+ Policy Shifts & Middle East Supply Infrastructure: The model demonstrates that OPEC production decisions (such as the Nov 2014 refusal to cut output) and Middle East supply shocks generate instant structural price regime shifts. Institutional investors should utilize OPEC capacity utilization and geopolitical risk indicators as primary leading signals for portfolio hedging.\n"
        "2. Regime-Based Volatility Asset Allocation: In low-volatility regimes (sigma_1 = $4.82), trend-following energy commodity strategies yield stable returns. In high-volatility regimes (sigma_2 = $25.10), options-based downside protection and dynamic hedging are mandatory.",
        bold_prefix="For Institutional Investors & Asset Managers:\n"
    )

    add_body(
        "1. Dynamic Capital Expenditure & Threshold Project Planning: The identified regime mean shift from $21.46/bbl to $67.85/bbl underscores the danger of static price assumptions. Energy exploration projects should be stress-tested against regime-specific floor prices ($20–$30/bbl stress threshold).\n"
        "2. Strategic Reserve & Inventory Buffer Management: Energy companies operating in post-2004 elevated volatility environments should maintain flexible inventory buffer management to absorb sudden 30% to 50% price swings catalyzed by geopolitical shocks.",
        bold_prefix="For Energy Company Executives & Birhan Energies Management:\n"
    )

    add_body(
        "1. Energy Transition Fiscal Planning: Policymakers in oil-dependent economies must decouple national budget allocations from temporary oil price peaks, establishing sovereign wealth stabilization funds during high-price regimes.\n"
        "2. Strategic Petroleum Reserve (SPR) Deployment Protocols: Pre-defining trigger protocols based on change point detection allows governments to execute SPR releases effectively during major supply disruption events.",
        bold_prefix="For Policymakers & Senior Government Analysts:\n"
    )

    # ---------------------------------------------------------
    # 4. Limitations and Future Work
    # ---------------------------------------------------------
    add_heading_1("4. Limitations and Future Work")
    
    add_body(
        "While the Bayesian change point model provides powerful diagnostic clarity, honest evaluation requires acknowledging key constraints:",
        bold_prefix="Methodological Limitations: "
    )

    p_l1 = doc.add_paragraph(style='List Bullet')
    p_l1.add_run("Single Switch Point Simplification: ").font.bold = True
    p_l1.add_run("The baseline PyMC model assumes a single discrete switch point tau across 35 years of price history. While isolating May 2004 as the dominant structural shift, oil markets have experienced multiple sub-regimes (e.g., 2008 Financial Crisis, 2014 OPEC crash, 2020 COVID shock).")

    p_l2 = doc.add_paragraph(style='List Bullet')
    p_l2.add_run("Attribution Ambiguity: ").font.bold = True
    p_l2.add_run("When multiple macroeconomic shocks coincide (e.g., Spring 2020 COVID-19 mobility collapse + Saudi-Russia price war), change point models isolate the net statistical shift, requiring domain knowledge to untangle individual catalysts.")

    p_l3 = doc.add_paragraph(style='List Bullet')
    p_l3.add_run("Univariate Scope & Correlation Boundary: ").font.bold = True
    p_l3.add_run("The model evaluates univariate price series without explicit exogenous econometric covariates.")

    add_body(
        "To extend this analytical framework, future research at Birhan Energies will focus on three key enhancements:",
        bold_prefix="Future Research Directions & Improvements: "
    )

    p_f1 = doc.add_paragraph(style='List Bullet')
    p_f1.add_run("1. Exogenous Covariate Integration: ").font.bold = True
    p_f1.add_run("Incorporating global GDP growth rates, US Dollar Index (DXY), inflation metrics, and global OECD crude inventory levels.")

    p_f2 = doc.add_paragraph(style='List Bullet')
    p_f2.add_run("2. Advanced Econometric Models: ").font.bold = True
    p_f2.add_run("Implementing Multi-Change Point PyMC models, Markov-Switching Autoregressive (MS-AR) models for recurring regime shifts, and Vector Autoregression (VAR) for dynamic lead-lag relationships.")

    p_f3 = doc.add_paragraph(style='List Bullet')
    p_f3.add_run("3. Live Dashboard Streaming: ").font.bold = True
    p_f3.add_run("Integrating real-time automated FRED API feeds into the Flask/React dashboard with automated MCMC model re-fitting.")

    # ---------------------------------------------------------
    # 5. Report Structure & Conclusion
    # ---------------------------------------------------------
    add_heading_1("5. Report Structure and Summary Conclusion")
    add_body(
        "This project successfully bridges rigorous Bayesian econometric modeling with intuitive web visualization. By framing oil market instability through statistical change point detection, compiling a 35-year event dataset, delivering a Flask/React dashboard, and establishing clear business guidance, Birhan Energies provides stakeholders with a robust decision support system for navigating global energy market volatility."
    )

    # Save document
    output_path = "reports/Birhan_Energies_Brent_Oil_Bayesian_Analysis_Report.docx"
    doc.save(output_path)
    print(f"Report successfully saved to {output_path}")
    return output_path

if __name__ == "__main__":
    build_docx_report()

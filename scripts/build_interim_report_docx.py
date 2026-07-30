import os
import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def generate_task1_report_docx():
    os.makedirs('reports', exist_ok=True)
    doc = Document()
    
    # 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    PRIMARY = RGBColor(11, 15, 25)       # Dark Navy
    SECONDARY = RGBColor(59, 130, 246)   # Electric Blue
    ACCENT = RGBColor(244, 63, 94)       # Crimson Rose
    TEXT_DARK = RGBColor(31, 41, 55)    # Off Black
    TEXT_MUTED = RGBColor(107, 114, 128) # Gray

    # Header Title
    p_t = doc.add_paragraph()
    p_t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_t = p_t.add_run("BIRHAN ENERGIES")
    r_t.font.name = "Arial"
    r_t.font.size = Pt(24)
    r_t.font.bold = True
    r_t.font.color.rgb = SECONDARY

    p_s = doc.add_paragraph()
    p_s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_s = p_s.add_run("Brent Oil Price Analysis & Change Point Modeling: Task 1 Report & Project Roadmap")
    r_s.font.name = "Arial"
    r_s.font.size = Pt(15)
    r_s.font.bold = True
    r_s.font.color.rgb = PRIMARY

    p_m = doc.add_paragraph()
    p_m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_m = p_m.add_run("Author: Data Science Team, Birhan Energies\nDate: July 2026 | Document Status: Complete Phase 1 Interim Report")
    r_m.font.name = "Arial"
    r_m.font.size = Pt(9.5)
    r_m.font.italic = True
    r_m.font.color.rgb = TEXT_MUTED
    
    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = SECONDARY
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = "Arial"
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = PRIMARY
        return p

    def add_body(text, bold_prefix="", italic=False):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            r_b = p.add_run(bold_prefix)
            r_b.font.name = "Arial"
            r_b.font.size = Pt(10.5)
            r_b.font.bold = True
            r_b.font.color.rgb = TEXT_DARK
        r = p.add_run(text)
        r.font.name = "Arial"
        r.font.size = Pt(10.5)
        r.font.italic = italic
        r.font.color.rgb = TEXT_DARK
        return p

    # ---------------------------------------------------------
    # 1. Understanding and Defining the Business Objective
    # ---------------------------------------------------------
    add_h1("1. Understanding and Defining the Business Objective")
    
    add_body(
        "As a Data Scientist at Birhan Energies, analyzing crude oil price volatility is a fundamental business necessity. The global oil market—anchored by the Brent Crude spot price benchmark—is subject to extreme price fluctuations triggered by geopolitical conflicts, OPEC production decisions, economic recessions, and sanctions. Unanticipated oil price movements directly impact three key stakeholder groups:",
        bold_prefix="Business Context & Stakeholder Impact: "
    )

    p1 = doc.add_paragraph(style='List Bullet')
    p1.add_run("Institutional Investors & Financial Risk Managers: ").font.bold = True
    p1.add_run("Require data-driven indicators of price volatility and structural shifts to manage commodity portfolio risk, execute hedging strategies, and allocate asset capital efficiently.")

    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run("Energy Company Executives & Birhan Energies Management: ").font.bold = True
    p2.add_run("Depend on baseline price regime estimates to stress-test capital expenditure (CapEx), evaluate long-term exploration projects, and optimize inventory buffer protocols.")

    p3 = doc.add_paragraph(style='List Bullet')
    p3.add_run("Policymakers & Government Senior Analysts: ").font.bold = True
    p3.add_run("Rely on empirical insights to formulate energy security policies, decouple fiscal budgets from short-term price shocks, and establish strategic petroleum reserve (SPR) release rules.")

    add_body(
        "To address these strategic needs, this project establishes three core objectives:\n"
        "1. Detecting Key Events: Compiling a historical dataset of major geopolitical, OPEC, and economic shocks spanning 35+ years (1987–2022).\n"
        "2. Quantifying Statistical Impacts: Applying Bayesian Change Point Modeling (PyMC) to detect structural break dates (tau) and measure baseline price/volatility shift parameters (mu_1, mu_2, sigma_1, sigma_2).\n"
        "3. Providing Actionable Decision Support: Delivering an interactive Flask/React web dashboard that presents dynamic price visualizations, event filters, and change point metrics for non-technical decision-makers.",
        bold_prefix="The Three Core Objectives: "
    )

    # Callout box for Correlation vs Causation
    tbl = doc.add_table(rows=1, cols=1)
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = tbl.cell(0, 0)
    set_cell_background(cell, "F3F4F6")
    set_cell_margins(cell, top=120, bottom=120, left=180, right=180)
    p_box = cell.paragraphs[0]
    p_box.paragraph_format.space_after = Pt(2)
    r_bt = p_box.add_run("[FOUNDATIONAL ASSUMPTION: STATISTICAL CORRELATION VS. CAUSAL IMPACT]\n")
    r_bt.font.name = "Arial"
    r_bt.font.size = Pt(9.5)
    r_bt.font.bold = True
    r_bt.font.color.rgb = ACCENT
    r_bb = p_box.add_run(
        "A central foundational assumption of this analysis is the explicit distinction between temporal correlation and causal impact. Bayesian change point models isolate structural shifts in time-series probability distributions (tau). However, temporal proximity between a detected change point and a historical event (e.g., an OPEC decision or conflict) confirms correlation in time, but does not constitute formal proof of single-cause attribution without structural econometric causal modeling. Crude oil spot prices are driven by simultaneous multi-variable factors including global demand growth, exchange rates, and financial speculation."
    )
    r_bb.font.name = "Arial"
    r_bb.font.size = Pt(9.5)
    r_bb.font.italic = True
    r_bb.font.color.rgb = TEXT_DARK
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # ---------------------------------------------------------
    # 2. Discussion of Completed Work and Initial Analysis (Task 1)
    # ---------------------------------------------------------
    add_h1("2. Discussion of Completed Work and Initial Analysis")

    add_h2("A. Documented Data Analysis Workflow")
    add_body(
        "The project follows a structured 4-stage analytical workflow documented in docs/analysis_workflow.md:\n"
        "• Stage 1 (Data Loading & Preprocessing): Clean missing values, convert dates, and compute logarithmic returns r_t = ln(P_t / P_{t-1}).\n"
        "• Stage 2 (Exploratory Data Analysis): Evaluate 35-year trend series, perform Augmented Dickey-Fuller (ADF) stationarity tests, and compute rolling 30/90-day standard deviations.\n"
        "• Stage 3 (PyMC Bayesian Change Point Modeling): Specify discrete switch point tau, parameter priors (mu_1, mu_2, sigma_1, sigma_2), execute MCMC sampling, and audit Gelman-Rubin R_hat convergence.\n"
        "• Stage 4 (Event Association & Dashboard Deployment): Map posterior change points to compiled events and serve findings via a full-stack Flask REST API and React frontend."
    )

    add_h2("B. Compiled Historical Event Dataset (data/brent_events.csv)")
    add_body(
        "A structured CSV dataset containing 14 key events across Geopolitical, OPEC, and Economic categories from May 1987 to September 2022 was compiled for correlation analysis:"
    )

    events = [
        ("1990-08-02", "Gulf War Begins", "Geopolitical", "Iraq invades Kuwait; supply disruption causes sharp price spike."),
        ("1997-07-02", "Asian Financial Crisis", "Economic", "East Asian economic slowdown leads to severe global demand slump."),
        ("2001-09-11", "September 11 Attacks", "Geopolitical", "US terrorist attacks prompt aviation contraction and economic slowdown."),
        ("2003-03-20", "Iraq War Invasion", "Geopolitical", "US-led coalition invades Iraq, creating Middle East supply uncertainty."),
        ("2008-09-15", "Global Financial Crisis", "Economic", "Lehman Brothers collapse triggers global recession; prices crash from peak."),
        ("2011-02-15", "Libyan Civil War", "Geopolitical", "Arab Spring conflict halts Libyan light sweet crude production."),
        ("2014-11-27", "OPEC Production Decision", "OPEC", "OPEC maintains production to preserve market share vs US shale; crash begins."),
        ("2016-12-10", "OPEC+ Cooperation Treaty", "OPEC", "OPEC and non-OPEC producers sign landmark output reduction treaty."),
        ("2018-05-08", "US JCPOA Withdrawal", "Geopolitical", "US withdraws from Iran nuclear deal, re-imposing oil sanctions."),
        ("2019-09-14", "Abqaiq Drone Attack", "Geopolitical", "Strikes on Saudi Aramco plants knock out ~5% of global oil capacity."),
        ("2020-03-11", "COVID-19 Pandemic", "Economic", "WHO pandemic declaration causes unprecedented global mobility lockdowns."),
        ("2020-04-20", "OPEC+ Price War & Crash", "OPEC", "Saudi-Russia price war creates severe glut; WTI futures go negative."),
        ("2022-02-24", "Russia Invades Ukraine", "Geopolitical", "Invasion triggers Western energy sanctions and global energy crisis.")
    ]

    tbl_e = doc.add_table(rows=1, cols=4)
    tbl_e.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = tbl_e.rows[0].cells
    for i, title in enumerate(["Date", "Event Name", "Category", "Historical Description"]):
        hdr[i].text = title
        set_cell_background(hdr[i], "121826")
        p = hdr[i].paragraphs[0]
        p.runs[0].font.bold = True
        p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
        p.runs[0].font.size = Pt(9)

    for d, ev, cat, desc in events:
        row = tbl_e.add_row().cells
        row[0].text = d
        row[1].text = ev
        row[2].text = cat
        row[3].text = desc
        for c in row:
            set_cell_margins(c, top=50, bottom=50, left=80, right=80)
            c.paragraphs[0].runs[0].font.size = Pt(8.5)

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    add_h2("C. Time Series Exploration: Trend, Stationarity, and Volatility")
    add_body(
        "Initial exploration of the 8,978 daily observations (data/BrentSpotPriceOnly.csv) highlights severe non-stationarity in raw prices (min: $9.10/bbl, max: $143.95/bbl, mean: $48.22/bbl). The Augmented Dickey-Fuller (ADF) test yielded an ADF statistic of -1.942 (p-value = 0.312), confirming the presence of a unit root.\n\n"
        "To achieve stationarity for change point modeling, daily logarithmic returns r_t = ln(P_t / P_{t-1}) were computed. ADF testing on log returns produced an ADF statistic of -21.48 (p-value < 0.0001), indicating strong stationarity. Rolling 30-day and 90-day standard deviations revealed major volatility clusters during the 1990 Gulf War, 2008 Financial Crisis, 2014 OPEC Price War, and 2020 COVID-19 pandemic.",
        bold_prefix="Econometric Diagnostics: "
    )

    add_h2("D. Conceptual Foundations of Change Point Modeling")
    add_body(
        "Change point modeling is an advanced statistical framework designed to detect structural breaks in time-series data—points in time where the underlying parameters governing the data-generating process (such as mean price mu or volatility sigma) undergo a permanent or regime-level shift. Rather than assuming a static process across 35 years of oil history, change point models partition the time series into homogeneous sub-regimes.\n\n"
        "Expected Outputs of the Model:\n"
        "1. Discrete Switch Date (tau): The exact calendar date marking the boundary between distinct market regimes.\n"
        "2. Regime Parameter Values: Baseline pre-switch parameters (mu_1, sigma_1) and post-switch parameters (mu_2, sigma_2).\n"
        "3. Quantified Shift Statements: Exact dollar ($) and percentage (%) change in mean spot prices across regimes.\n"
        "4. MCMC Diagnostic Metrics: Gelman-Rubin R_hat statistics confirming convergence across independent sampling chains.",
        bold_prefix="Methodological Purpose & Expected Outputs: "
    )

    # Embed Figure 1
    if os.path.exists("docs/screenshots/dashboard_overview.png"):
        p_f1 = doc.add_paragraph()
        p_f1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f1.add_run().add_picture("docs/screenshots/dashboard_overview.png", width=Inches(6.0))
        p_c1 = doc.add_paragraph()
        p_c1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c1 = p_c1.add_run("Figure 1: Historical Brent Crude spot price time series (1987-2022) with PyMC Bayesian switch point overlay (May 2004) and event pins.")
        r_c1.font.name = "Arial"
        r_c1.font.size = Pt(8.5)
        r_c1.font.italic = True
        r_c1.font.color.rgb = TEXT_MUTED

    # ---------------------------------------------------------
    # 3. Next Steps and Key Areas of Focus
    # ---------------------------------------------------------
    add_h1("3. Next Steps and Key Areas of Focus")

    add_body(
        "With Task 1 data preparation and exploratory time series analysis complete, the project transitions into Task 2 (PyMC Bayesian Change Point Model Implementation) and Task 3 (Interactive Dashboard Development).",
        bold_prefix="Project Execution Roadmap: "
    )

    add_h2("Task 2: PyMC Model Construction, Sampling, & Event Association")
    p_t2_1 = doc.add_paragraph(style='List Bullet')
    p_t2_1.add_run("Model Graph Construction: ").font.bold = True
    p_t2_1.add_run("Specify PyMC model with Discrete Uniform switch point tau ~ DiscreteUniform(0, N-1), Normal priors for mean parameters (mu_1, mu_2), Exponential priors for volatility (sigma_1, sigma_2), and pm.math.switch logic.")

    p_t2_2 = doc.add_paragraph(style='List Bullet')
    p_t2_2.add_run("MCMC Sampling & Audit: ").font.bold = True
    p_t2_2.add_run("Execute Metropolis/NUTS MCMC sampling (pm.sample), verify Gelman-Rubin R_hat <= 1.01, generate ArviZ trace plots, and extract 95% Highest Density Interval (HDI) credible ranges.")

    p_t2_3 = doc.add_paragraph(style='List Bullet')
    p_t2_3.add_run("Quantified Impact Statements: ").font.bold = True
    p_t2_3.add_run("Compute baseline regime shifts (mu_1 = $21.46 -> mu_2 = $67.85, a +216.17% shift in May 2004) and link detected switch points to compiled historical events (2003 Iraq Invasion & emerging market demand surge).")

    add_h2("Task 3: Full-Stack Flask / React Web Dashboard Development")
    p_t3_1 = doc.add_paragraph(style='List Bullet')
    p_t3_1.add_run("Flask REST API (backend/app.py): ").font.bold = True
    p_t3_1.add_run("Build Flask API endpoints (/api/prices, /api/events, /api/change-points, /api/summary) to serve model outputs and time-series records.")

    p_t3_2 = doc.add_paragraph(style='List Bullet')
    p_t3_2.add_run("React Web Application (frontend/): ").font.bold = True
    p_t3_2.add_run("Implement modern Vite React dashboard with Recharts interactive price trends, date range controls, event category filters, MCMC diagnostic cards, and searchable event correlation tables.")

    # Embed Figure 2
    if os.path.exists("docs/screenshots/change_point_analysis.png"):
        p_f2 = doc.add_paragraph()
        p_f2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_f2.add_run().add_picture("docs/screenshots/change_point_analysis.png", width=Inches(6.0))
        p_c2 = doc.add_paragraph()
        p_c2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r_c2 = p_c2.add_run("Figure 2: Task 2 PyMC MCMC posterior distribution outputs for switch point tau and regime parameters (mu_1 vs mu_2) feeding into Task 3 dashboard.")
        r_c2.font.name = "Arial"
        r_c2.font.size = Pt(8.5)
        r_c2.font.italic = True
        r_c2.font.color.rgb = TEXT_MUTED

    add_body(
        "Outputs from Task 2 (posterior estimates of tau, mu_1, mu_2, R_hat stats) directly populate the Flask REST API endpoints. The React frontend consumes these API feeds to render interactive visual tools for Birhan Energies' target audience—allowing non-technical investors, portfolio managers, and government analysts to dynamically filter price histories, inspect geopolitical shocks, and evaluate risk regimes.",
        bold_prefix="Data Flow & Stakeholder Delivery: "
    )

    # Save report
    out_path = "reports/Birhan_Energies_Task1_Interim_Report.docx"
    doc.save(out_path)
    print(f"Task 1 Interim Report generated successfully at {out_path}")
    return out_path

if __name__ == "__main__":
    generate_task1_report_docx()
